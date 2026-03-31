from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import RGBColor, Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import os
import requests
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)


# ------------------------------------------------------------
# BASE TEMPLATE (optional)
# ------------------------------------------------------------

BASE_DIR = os.path.dirname(__file__)
BASE_TEMPLATE = os.path.join(BASE_DIR, "base.docx")


# ------------------------------------------------------------
# GLOBAL TYPOGRAPHY CONFIG (MATCHES EDITOR / globals.css)
# ------------------------------------------------------------

BODY_FONT   = "Times New Roman"
BODY_SIZE   = 12

H1_SIZE     = 22     # "Statement of Work" — big, centered, uppercase
H2_SIZE     = 13     # Section headings — left-border accent
H3_SIZE     = 11.5

# Brand blue that matches --primary: #2563eb
H2_BORDER_COLOR = "2563EB"   # left accent border on h2
H2_BORDER_SIZE  = 24         # 1/8-pt units → 3 pt thick

META_LABEL_FILL  = "EFF6FF"  # --muted (blue-50)
META_HEADER_FILL = "F8FAFC"  # table header background


# ------------------------------------------------------------
# STYLE CONFIGURATION
# ------------------------------------------------------------

def configure_document_styles(document: Document):
    styles = document.styles

    try:
        normal = styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(BODY_SIZE)
    except KeyError:
        pass

    heading_map = {
        "Heading 1": H1_SIZE,
        "Heading 2": H2_SIZE,
        "Heading 3": H3_SIZE,
    }
    for style_name, size in heading_map.items():
        try:
            s = styles[style_name]
            s.font.name = BODY_FONT
            s.font.size = Pt(size)
            s.font.bold = True
        except KeyError:
            continue


# ------------------------------------------------------------
# PARAGRAPH SPACING
# ------------------------------------------------------------

def apply_body_spacing(paragraph):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after  = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.35


# ------------------------------------------------------------
# H2 LEFT ACCENT BORDER  (replicates CSS border-left: 4px solid var(--primary))
# ------------------------------------------------------------

def add_left_border(paragraph, color=H2_BORDER_COLOR, size=H2_BORDER_SIZE):
    """
    Adds a left paragraph border to mimic the editor's
      border-left: 4px solid var(--primary)
    on h2 headings.
    """
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")

    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(size))   # 1/8-pt → 24 = 3pt
    left.set(qn("w:space"), "12")        # gap between border and text (twips)
    left.set(qn("w:color"), color)

    pBdr.append(left)
    pPr.append(pBdr)


# ------------------------------------------------------------
# CELL SHADING
# ------------------------------------------------------------

def shade_cell(cell, fill_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


# ------------------------------------------------------------
# CELL BORDER HELPERS
# ------------------------------------------------------------

def remove_cell_borders(cell):
    """Remove all borders from a cell (used for signature table)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_bottom_border_to_cell(cell, color="334155", size=6):
    """Add only a bottom border (signature line)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(size))
    bottom.set(qn("w:color"), color)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


# ------------------------------------------------------------
# TEXT RUNS  (handles bold, italic, underline, color)
# Note: formyxaField nodes are converted to text in route.js
#       before reaching Flask. This handles any that slip through.
# ------------------------------------------------------------

def add_text_runs_from_tiptap(content_nodes, paragraph):
    if not content_nodes:
        return

    for node in content_nodes:
        ntype = node.get("type")

        # ── Standard text node ──────────────────────────────────────────────
        if ntype == "text":
            text  = node.get("text", "")
            marks = node.get("marks", []) or []

            run = paragraph.add_run(text)
            run.font.name = BODY_FONT
            run.font.size = Pt(BODY_SIZE)

            for m in marks:
                mtype = m.get("type")
                if mtype == "bold":      run.bold      = True
                if mtype == "italic":    run.italic    = True
                if mtype == "underline": run.underline = True

            for m in marks:
                if m.get("type") == "textStyle":
                    color = m.get("attrs", {}).get("color", "")
                    if color and color.startswith("#") and len(color) == 7:
                        try:
                            r = int(color[1:3], 16)
                            g = int(color[3:5], 16)
                            b = int(color[5:7], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
                        except Exception:
                            pass

        # ── formyxaField fallback (should already be converted in route.js) ─
        elif ntype == "formyxaField":
            attrs  = node.get("attrs", {}) or {}
            value  = (attrs.get("value") or "").strip()
            label  = (attrs.get("label") or "Field").strip()
            is_bold = bool(attrs.get("bold"))

            display = value if value else f"[{label}]"
            run = paragraph.add_run(display)
            run.font.name = BODY_FONT
            run.font.size = Pt(BODY_SIZE)
            if is_bold:
                run.bold = True
            if not value:
                # Dotted underline for unfilled placeholders
                run.underline = True


# ------------------------------------------------------------
# META-TABLE  (the Client / Service Provider / Effective Date grid)
# Detected via node.attrs.class == "meta-table"
# ------------------------------------------------------------

def render_meta_table(node, document: Document):
    """
    Renders the meta-grid table at the top of the SOW with the
    correct label/value styling that mirrors the CSS .meta-table rules:
      - Label cells: uppercase, small, bold, blue-50 background
      - Value cells: normal weight, white background
    """
    rows = node.get("content", [])
    if not rows:
        return

    num_rows = len(rows)
    num_cols = max(len(r.get("content", [])) for r in rows)

    table = document.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit   = True

    for r_idx, row in enumerate(rows):
        cells = row.get("content", [])
        for c_idx, cell_node in enumerate(cells):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""

            for child in cell_node.get("content", []):
                if child.get("type") != "paragraph":
                    continue
                content = child.get("content", []) or []

                p = cell.paragraphs[0]
                add_text_runs_from_tiptap(content, p)
                fmt = p.paragraph_format
                fmt.space_before = Pt(2)
                fmt.space_after  = Pt(2)

            # ── Style: label cols are odd-indexed (0, 2), value cols are even (1, 3) ──
            is_label_col = (c_idx % 2 == 0)
            is_header_cell = (cell_node.get("type") == "tableHeader")

            if is_label_col or is_header_cell:
                shade_cell(cell, META_LABEL_FILL)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.name = BODY_FONT
                        # Simulate text-transform: uppercase
                        run.text = run.text.upper()
            else:
                shade_cell(cell, "FFFFFF")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(BODY_SIZE)
                        run.font.name = BODY_FONT

    # Add spacing after table
    document.add_paragraph()


# ------------------------------------------------------------
# STANDARD TABLE
# ------------------------------------------------------------

def render_table_node(node, document: Document):
    rows = node.get("content", [])
    if not rows:
        return

    num_rows = len(rows)
    num_cols = max(len(r.get("content", [])) for r in rows)

    table = document.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit   = True

    for r_idx, row in enumerate(rows):
        cells = row.get("content", [])
        for c_idx, cell_node in enumerate(cells):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""

            for child in cell_node.get("content", []):
                if child.get("type") != "paragraph":
                    continue

                attrs = child.get("attrs", {}) or {}
                if attrs.get("instructional") is True:
                    continue

                content = child.get("content", []) or []
                if not any(
                    c.get("type") == "text" and c.get("text", "").strip()
                    for c in content
                ):
                    continue

                p = cell.paragraphs[0]
                add_text_runs_from_tiptap(content, p)
                apply_body_spacing(p)

            # Header styling
            if cell_node.get("type") == "tableHeader":
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                shade_cell(cell, META_HEADER_FILL)


# ------------------------------------------------------------
# SIGNATURES BLOCK
# Replicates the signaturesBlock node with two signature boxes
# ------------------------------------------------------------

def render_signatures_block(node, document: Document):
    attrs       = node.get("attrs", {}) or {}
    left_title  = attrs.get("leftTitle",  "CLIENT")
    right_title = attrs.get("rightTitle", "SERVICE PROVIDER")

    # Spacer before signatures
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(24)
    spacer.paragraph_format.space_after  = Pt(6)

    # 2-column table, no visible outer borders — just two signature blocks
    table = document.add_table(rows=5, cols=2)
    table.alignment  = WD_TABLE_ALIGNMENT.CENTER
    table.autofit    = True

    # Remove all cell borders first, then selectively add bottom lines
    for row in table.rows:
        for cell in row.cells:
            remove_cell_borders(cell)

    def set_cell_text(cell, text, bold=False, size=BODY_SIZE):
        cell.text = ""
        p   = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = BODY_FONT
        run.font.size = Pt(size)
        if bold:
            run.bold = True
        fmt = p.paragraph_format
        fmt.space_before = Pt(2)
        fmt.space_after  = Pt(4)

    # Row 0: Party titles (CLIENT / SERVICE PROVIDER)
    set_cell_text(table.rows[0].cells[0], left_title,  bold=True, size=10)
    set_cell_text(table.rows[0].cells[1], right_title, bold=True, size=10)

    # Row 1: Signature line (empty cell with bottom border = signature line)
    for cell in table.rows[1].cells:
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(24)  # vertical space for signature
        p.paragraph_format.space_after  = Pt(2)
        add_bottom_border_to_cell(cell)

    # Row 2: "Signature" label
    set_cell_text(table.rows[2].cells[0], "Signature", size=9)
    set_cell_text(table.rows[2].cells[1], "Signature", size=9)

    # Row 3: Name line
    for cell in table.rows[3].cells:
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(2)
        add_bottom_border_to_cell(cell)

    # Row 4: Name + Date labels
    set_cell_text(table.rows[4].cells[0], "Name / Date", size=9)
    set_cell_text(table.rows[4].cells[1], "Name / Date", size=9)


# ------------------------------------------------------------
# IMAGE SUPPORT
# ------------------------------------------------------------

def render_image_node(node, document: Document):
    attrs = node.get("attrs", {})
    src   = attrs.get("src")
    if not src:
        return
    try:
        resp = requests.get(src, timeout=8)
        resp.raise_for_status()
        image_bytes = io.BytesIO(resp.content)
    except Exception:
        return

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_bytes, width=Inches(4.5))
    apply_body_spacing(paragraph)


# ------------------------------------------------------------
# CORE RENDER FUNCTION
# ------------------------------------------------------------

def render_node(node, document: Document):
    ntype = node.get("type")

    # ── Heading ─────────────────────────────────────────────────────────────
    if ntype == "heading":
        level = node.get("attrs", {}).get("level", 1)
        p     = document.add_paragraph()

        size_map = {1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE}
        size = size_map.get(level, BODY_SIZE)

        add_text_runs_from_tiptap(node.get("content", []), p)

        for run in p.runs:
            run.bold      = True
            run.font.size = Pt(size)
            run.font.name = BODY_FONT

            # H1: uppercase to match CSS text-transform: uppercase
            if level == 1:
                run.text = run.text.upper()

        fmt = p.paragraph_format
        if level == 1:
            # Centered, spacious — matches editor h1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_before = Pt(14)
            fmt.space_after  = Pt(20)
            fmt.line_spacing = 1.3

        elif level == 2:
            # Left-aligned with left accent border — mirrors CSS border-left
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.space_before = Pt(18)
            fmt.space_after  = Pt(6)
            fmt.line_spacing = 1.3
            add_left_border(p)

        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.space_before = Pt(12)
            fmt.space_after  = Pt(4)
            fmt.line_spacing = 1.3

        return

    # ── Paragraph ───────────────────────────────────────────────────────────
    if ntype == "paragraph":
        attrs = node.get("attrs", {}) or {}

        # Skip instructional ghost paragraphs
        if attrs.get("instructional") is True:
            return

        content = node.get("content", []) or []
        if not any(
            c.get("type") == "text" and c.get("text", "").strip()
            for c in content
        ):
            return

        p = document.add_paragraph()
        add_text_runs_from_tiptap(content, p)
        apply_body_spacing(p)
        return

    # ── Bullet list ──────────────────────────────────────────────────────────
    if ntype == "bulletList":
        for li in node.get("content", []):
            if li.get("type") != "listItem":
                continue
            for child in li.get("content", []):
                if child.get("type") != "paragraph":
                    continue

                attrs = child.get("attrs", {}) or {}
                if attrs.get("instructional") is True:
                    continue

                content = child.get("content", []) or []
                if not any(
                    c.get("type") == "text" and c.get("text", "").strip()
                    for c in content
                ):
                    continue

                p = document.add_paragraph()
                bullet_run = p.add_run("•  ")
                bullet_run.font.name = BODY_FONT
                bullet_run.font.size = Pt(BODY_SIZE)
                add_text_runs_from_tiptap(content, p)
                fmt = p.paragraph_format
                fmt.left_indent   = Pt(12)
                fmt.space_before  = Pt(0)
                fmt.space_after   = Pt(5)
                fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                fmt.line_spacing  = 1.35
        return

    # ── Ordered list ─────────────────────────────────────────────────────────
    if ntype == "orderedList":
        index = 1
        for li in node.get("content", []):
            for child in li.get("content", []):
                if child.get("type") == "paragraph":
                    p = document.add_paragraph()
                    num_run = p.add_run(f"{index}.  ")
                    num_run.font.name = BODY_FONT
                    num_run.font.size = Pt(BODY_SIZE)
                    add_text_runs_from_tiptap(child.get("content", []), p)
                    apply_body_spacing(p)
                    index += 1
        return

    # ── Table (with meta-table detection) ────────────────────────────────────
    if ntype == "table":
        table_class = (node.get("attrs") or {}).get("class", "")
        if table_class == "meta-table":
            render_meta_table(node, document)
        else:
            render_table_node(node, document)
        return

    # ── Signatures block ──────────────────────────────────────────────────────
    if ntype == "signaturesBlock":
        render_signatures_block(node, document)
        return

    # ── Images ───────────────────────────────────────────────────────────────
    if ntype in ("image", "resizableImage"):
        render_image_node(node, document)
        return

    # ── Silently skip unknown nodes (coverMetadataBlock, etc.) ───────────────
    # No crash, no content — clean skip.
    return


# ------------------------------------------------------------
# MAIN CONVERTER
# ------------------------------------------------------------

def tiptap_doc_to_docx(tiptap_doc):
    if os.path.exists(BASE_TEMPLATE):
        document = Document(BASE_TEMPLATE)
    else:
        document = Document()

    configure_document_styles(document)

    if not tiptap_doc or tiptap_doc.get("type") != "doc":
        return document

    for node in tiptap_doc.get("content", []):
        render_node(node, document)

    return document


# ------------------------------------------------------------
# ENDPOINT
# ------------------------------------------------------------

@app.route("/generate-docx", methods=["POST"])
def generate_docx():
    try:
        data      = request.get_json(force=True)
        tiptap_doc = data.get("contentJson")
        file_name  = data.get("fileName", "document.docx")

        document = tiptap_doc_to_docx(tiptap_doc)

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        return send_file(
            buf,
            as_attachment=True,
            download_name=file_name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8001, debug=True)